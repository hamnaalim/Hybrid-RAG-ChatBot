"""Document loaders for PDF, DOCX, TXT, and HTML files."""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".html", ".htm", ".md"}


@dataclass
class Document:
    """A loaded document with text content and metadata."""

    content: str
    source: str
    metadata: dict = field(default_factory=dict)

    def __post_init__(self) -> None:
        self.metadata.setdefault("source", self.source)


def load_pdf(path: Path) -> list[Document]:
    """Load a PDF, one Document per page when page text is available."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    docs: list[Document] = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        docs.append(
            Document(
                content=text,
                source=path.name,
                metadata={
                    "source": path.name,
                    "path": str(path),
                    "page": i + 1,
                    "file_type": "pdf",
                },
            )
        )
    if not docs:
        raise ValueError(f"No extractable text found in PDF: {path}")
    return docs


def load_docx(path: Path) -> list[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    content = "\n\n".join(paragraphs)
    if not content:
        raise ValueError(f"No extractable text found in DOCX: {path}")
    return [
        Document(
            content=content,
            source=path.name,
            metadata={
                "source": path.name,
                "path": str(path),
                "page": 1,
                "file_type": "docx",
            },
        )
    ]


def load_txt(path: Path) -> list[Document]:
    content = path.read_text(encoding="utf-8", errors="ignore").strip()
    if not content:
        raise ValueError(f"Empty text file: {path}")
    return [
        Document(
            content=content,
            source=path.name,
            metadata={
                "source": path.name,
                "path": str(path),
                "page": 1,
                "file_type": path.suffix.lstrip(".").lower() or "txt",
            },
        )
    ]


def load_html(path: Path) -> list[Document]:
    from bs4 import BeautifulSoup

    raw = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    content = " ".join(soup.get_text(separator=" ").split())
    if not content:
        raise ValueError(f"No extractable text found in HTML: {path}")
    return [
        Document(
            content=content,
            source=path.name,
            metadata={
                "source": path.name,
                "path": str(path),
                "page": 1,
                "file_type": "html",
            },
        )
    ]


_LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".txt": load_txt,
    ".md": load_txt,
    ".html": load_html,
    ".htm": load_html,
}


def load_file(path: Path | str) -> list[Document]:
    """Load a single file into one or more Documents."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)
    ext = path.suffix.lower()
    loader = _LOADERS.get(ext)
    if loader is None:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {sorted(SUPPORTED_EXTENSIONS)}")
    logger.info("Loading %s", path)
    return loader(path)


def load_directory(directory: Path | str) -> list[Document]:
    """Recursively load all supported documents from a directory."""
    directory = Path(directory)
    if not directory.exists():
        raise FileNotFoundError(directory)

    documents: list[Document] = []
    files = sorted(
        p for p in directory.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
    )
    for path in files:
        try:
            documents.extend(load_file(path))
        except Exception as exc:  # noqa: BLE001 — keep ingesting other files
            logger.warning("Failed to load %s: %s", path, exc)
    logger.info("Loaded %d document parts from %d files", len(documents), len(files))
    return documents


def load_paths(paths: Iterable[Path | str]) -> list[Document]:
    documents: list[Document] = []
    for path in paths:
        documents.extend(load_file(path))
    return documents
